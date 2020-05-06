import docx
import os


li = os.listdir("E:/test")
for n in li:
    doc = docx.Document()
    out = "E:/demo/%s(封面）.docx"%n.strip(".txt")
    doc.add_heading(n.strip(".txt"),0)
    doc.save(out)

for n in li:
    file = "E:/test/%s"%n
    out = "E:/demo/%s.docx"%n.strip(".txt")
    file = open(file,"r")
    doc = docx.Document()
    for i in file.readlines():
        doc.add_paragraph(i)
    doc.save(out)

print("over")
